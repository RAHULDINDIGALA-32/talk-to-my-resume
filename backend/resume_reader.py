from pathlib import Path
from pypdf import PdfReader
from docx import Document

def read_pdf(file_path: str):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file: str):
    document  = Document(file)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path: str):
    file_extension = Path(file_path).suffix.lower()

    if file_extension == ".pdf":
        return read_pdf(file_path)
    elif file_extension == ".docx":
        return read_docx(file_path)
    else:
        return None